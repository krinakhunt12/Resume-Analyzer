"""
Text Extractor Module
Handles extraction of text from PDF and DOCX files
"""

import pdfplumber
from docx import Document
import os


class TextExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
    
    def extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    def extract(self, file_path):
        """
        Main extraction method that determines file type and extracts text
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted text from the document
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")
        
        if file_ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        
    def get_file_info(self, file_path):
        """Get basic file information"""
        if not os.path.exists(file_path):
            return None
        
        file_info = {
            'filename': os.path.basename(file_path),
            'size_kb': round(os.path.getsize(file_path) / 1024, 2),
            'extension': os.path.splitext(file_path)[1],
            'path': file_path
        }
        
        return file_info